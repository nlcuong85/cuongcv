from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.shared import Cm, Pt


SOURCE = Path("tmp/docs/Lohtor_application_form_source.docx")
OUTPUT = Path("output/doc/Lohtor.Nguyen.LeCuong.docx")
PHOTO = Path("public/images/avatar.jpg")


def add_value(paragraph, value, tab=True, size=10):
    """Append a compact, clearly handwritten-style field value."""
    run = paragraph.add_run(("\t" if tab else "") + value)
    run.bold = False
    run.font.size = Pt(size)
    return run


def replace_line(paragraph, text, size=10):
    for run in paragraph.runs:
        run.text = ""
    run = paragraph.add_run(text)
    run.bold = False
    run.font.size = Pt(size)
    return run


doc = Document(SOURCE)
p = doc.paragraphs

# Requested rental start: mark September 1; intentionally leave the other-period row blank.
p[2].runs[-1].text = " [X]\t\t\t\t\t\t"
photo_run = p[2].add_run()
photo_run.add_picture(str(PHOTO), width=Cm(2.0), height=Cm(2.0))
p[4].runs[-1].text = ""

# Personal and contact details. Unprovided sensitive fields remain blank.
for index, value in {
    6: "Nguyen",
    7: "Le Cuong",
    8: "Male [X] / female [ ]",
    12: "Vietnamese",
    15: "+49 163 3628505",
    17: "nlcuong999@gmail.com",
    18: "Gildenstrasse 13",
    19: "74074",
    20: "Heilbronn",
    21: "Germany",
}.items():
    add_value(p[index], value)

# Accommodation and services: single apartment, no TV/telephone, internet required.
replace_line(p[24], "- Apartment mit Bad / incl. bathroom                         yes [X]  no [ ]")
replace_line(p[25], "- Apartment Comfort mit Bad und Kochnische / incl. bathroom and kitchen    yes [ ]  no [X]")
replace_line(p[26], "- Apartment double fuer Paare mit Doppelbett / for couples incl. double bed  yes [ ]  no [X]")
replace_line(p[27], "- 2er-4er WG / 2-4 person shared apartment                    yes [ ]  no [X]")
replace_line(p[30], "- TV-connection                                             yes [ ]  no [X]")
replace_line(p[31], "- Telephone-connection                                      yes [ ]  no [X]")
replace_line(p[32], "- Internet-connection                                       yes [X]  no [ ]")
replace_line(p[33], "Haben Sie ein Auto? / Do you have a car?                    yes [ ]  no [X]")
replace_line(p[34], "Raucher / Smoker?                                           yes [ ]  no [X]  gelegentlich/sometimes [ ]")

# Study and eligibility information.
replace_line(p[36], "Name der Hochschule / Name of University: Heilbronn University")
replace_line(p[37], "Studiengang / Program name: Master in Software Engineering")
replace_line(p[39], "Im wievielten Semester studieren Sie im genannten Studiengang? / Current semester: 2")
replace_line(p[41], "Wieviele Semester moechten Sie im Haus wohnen? / Intended stay: 1 semester (possibly longer)")
replace_line(p[43], "Immatrikulationsbescheinigung vorhanden? / certificate of enrolment?  Yes [X]  No [ ]")

# Motivation and the requested work-student clarification.
replace_line(p[61], "Weshalb moechten Sie im Haus Lohtor wohnen? / Why do you want to live in house Lohtor?")
replace_line(p[62], "I am a member of Experimenta Maker Lab and would like to live nearby so I can work on my project there.")
replace_line(p[66], "Unternehmen/company: Working student (contract available)")
replace_line(p[67], "Universitaet/University: Heilbronn University")
replace_line(p[71], "Praktikumsvertrag vorhanden? / Do you have an internship work contract?  Yes [ ]  No [X] (working-student contract available)")

# Exchange-student information as specified by the user.
replace_line(p[74], "Name der Heimatuniversitaet / name of your home university: Vietnamese-German University (VGU)")
replace_line(p[75], "Stadt / City: Ho Chi Minh City")
replace_line(p[76], "Land / country: Vietnam")

OUTPUT.parent.mkdir(parents=True, exist_ok=True)
doc.save(OUTPUT)
print(OUTPUT)
